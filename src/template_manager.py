from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


BUSINESS_TRIP_TAXI_TEMPLATE = "Служебная_записка_такси.docx"


REPORT_TEMPLATE_FILES = {
    "business_trip": {
        BUSINESS_TRIP_TAXI_TEMPLATE: "Служебная записка о компенсации расходов на такси",
    },
    "representative_expenses": {
        "Смета_и_отчет_представительские.docx": "Смета и отчёт о представительских расходах",
    },
    "gifts": {
        "Служебная_записка_подарки.docx": "Служебная записка на приобретение подарков",
    },
}


class TemplateManager:
    def __init__(self, templates_dir: Path):
        self.templates_dir = templates_dir

    def ensure_default_templates(self, report_type: str | None = None) -> None:
        report_types = [report_type] if report_type else list(REPORT_TEMPLATE_FILES)
        for current_type in report_types:
            target_dir = self.templates_dir / current_type
            target_dir.mkdir(parents=True, exist_ok=True)
            for file_name, title in REPORT_TEMPLATE_FILES[current_type].items():
                path = target_dir / file_name
                if not path.exists():
                    create_default_template(path, title, current_type)

    def templates_for(self, report_type: str) -> list[Path]:
        self.ensure_default_templates(report_type)
        active_templates = set(REPORT_TEMPLATE_FILES[report_type])
        return sorted(
            path
            for path in (self.templates_dir / report_type).glob("*.docx")
            if path.name in active_templates
        )

    def save_uploaded_templates(self, report_type: str, uploaded_files: list) -> list[Path]:
        target_dir = self.templates_dir / report_type
        target_dir.mkdir(parents=True, exist_ok=True)
        saved: list[Path] = []
        for uploaded_file in uploaded_files:
            path = target_dir / uploaded_file.name
            path.write_bytes(uploaded_file.getvalue())
            saved.append(path)
        return saved


def create_default_template(path: Path, title: str, report_type: str) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    for style_name in ("Heading 1", "Heading 2"):
        style = styles[style_name]
        style.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    heading = document.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("Дата составления: {{ report_date }}")
    document.add_paragraph("Организация: {{ employee.company }}{{ initiator.company }}")
    document.add_paragraph("Сотрудник: {{ employee.full_name }}{{ initiator.full_name }}")
    document.add_paragraph("Должность: {{ employee.position }}{{ initiator.position }}")

    if report_type == "business_trip":
        _add_business_trip_body(document)
    elif report_type == "representative_expenses":
        _add_representative_body(document)
    elif report_type == "gifts":
        _add_gifts_body(document)

    document.add_heading("Сведения о чеках", level=2)
    table = document.add_table(rows=2, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["N", "Дата", "Продавец", "Тип", "Сумма"]
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for index, value in enumerate(["{{ receipts_table }}", "", "", "", ""]):
        table.rows[1].cells[index].text = value

    document.add_paragraph("Итого: {{ total_amount }} ({{ total_amount_words }})")
    document.add_paragraph("Подписант: {{ employee.default_signatory_name }}{{ initiator.default_signatory_name }}")
    document.save(str(path))


def _add_business_trip_body(document: Document) -> None:
    document.add_heading("Данные командировки", level=2)
    for text in [
        "Город командировки: {{ trip_city }}",
        "Период: {{ trip_start_date }} - {{ trip_end_date }}",
        "Цель поездки: {{ purpose }}",
        "Объект / проект: {{ project }}",
        "Маршрут: {{ route }}",
        "Контрагент: {{ counterparty }}",
        "Основание: {{ basis }}",
        "Комментарий: {{ comment }}",
    ]:
        document.add_paragraph(text)


def _add_representative_body(document: Document) -> None:
    document.add_heading("Данные мероприятия", level=2)
    for text in [
        "Дата мероприятия: {{ event_date }}",
        "Место проведения: {{ place }}",
        "Ресторан / кафе: {{ restaurant_name }}",
        "Контрагент: {{ counterparty }}",
        "Цель встречи: {{ meeting_purpose }}",
        "Участники компании: {{ participants_company_text }}",
        "Участники контрагента: {{ participants_counterparty_text }}",
        "Результат встречи: {{ meeting_result }}",
    ]:
        document.add_paragraph(text)


def _add_gifts_body(document: Document) -> None:
    document.add_heading("Данные по подаркам", level=2)
    for text in [
        "Дата покупки: {{ purchase_date }}",
        "Наименование подарков: {{ gift_name }}",
        "Количество: {{ gift_quantity }}",
        "Стоимость за единицу: {{ unit_price }}",
        "Получатели: {{ recipients_text }}",
        "Контрагент: {{ counterparty }}",
        "Повод: {{ occasion }}",
        "Цель расходов: {{ purpose }}",
    ]:
        document.add_paragraph(text)
