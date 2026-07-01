from __future__ import annotations

import re
from dataclasses import dataclass
from datetime import date
from decimal import Decimal
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt

from .approval import default_approver_for
from .docx_generator import DocxGenerator
from .formatters import amount_to_words_ru, format_date_ru, format_rubles
from .models import BusinessTripReport, Employee, GiftExpenseReport, Receipt, RepresentativeExpenseReport
from .template_manager import TemplateManager
from .utils import slugify_file_part, unique_path


DEFAULT_SIGNATORY_NAME = "Хуан Голян"
DEFAULT_SIGNATORY_NAME_DATIVE = "Хуан Голяну"
DEFAULT_SIGNATORY_POSITION = "Генеральному директору"


@dataclass
class BuildResult:
    files: list[Path]
    warnings: list[str]


class BaseReportBuilder:
    report_type: str

    def __init__(self, template_manager: TemplateManager, output_dir: Path, generator: DocxGenerator | None = None):
        self.template_manager = template_manager
        self.output_dir = output_dir
        self.generator = generator or DocxGenerator()

    def build(self, report: Any) -> BuildResult:
        context = self.build_context(report)
        files: list[Path] = []
        warnings: list[str] = []
        for template_path in self.template_manager.templates_for(self.report_type):
            output_path = self.output_dir / self.build_file_name(report, template_path)
            missing = self.generator.find_unfilled_placeholders(template_path, context)
            if missing:
                warnings.append(f"{template_path.name}: не заполнены {', '.join(sorted(missing))}")
            files.append(self.generator.render(template_path, context, output_path))
        return BuildResult(files=files, warnings=warnings)

    def build_context(self, report: Any) -> dict[str, Any]:
        raise NotImplementedError

    def build_file_name(self, report: Any, template_path: Path) -> str:
        raise NotImplementedError


class BusinessTripBuilder(BaseReportBuilder):
    report_type = "business_trip"

    def build(self, report: BusinessTripReport) -> BuildResult:
        result = super().build(report)
        for file_path in result.files:
            _compact_business_trip_taxi_memo(file_path, receipts_count=len(report.receipts))
        return result

    def build_context(self, report: BusinessTripReport) -> dict[str, Any]:
        context = _base_context(report.receipts, report.report_date, employee=report.employee)
        context["taxi_compensation_purpose"] = _business_trip_taxi_compensation_purpose(
            report.receipts,
            trip_city=report.trip_city,
            trip_start_date=report.trip_start_date,
            trip_end_date=report.trip_end_date,
            trip_purpose=report.purpose,
        )
        context["taxi_compensation_purpose_label"] = "Цель поездок" if len(report.receipts) > 1 else "Цель поездки"
        context.update(
            {
                "trip_city": report.trip_city,
                "trip_start_date": format_date_ru(report.trip_start_date),
                "trip_end_date": format_date_ru(report.trip_end_date),
                "purpose": report.purpose,
                "project": report.project or "",
                "route": report.route or "",
                "counterparty": report.counterparty or "",
                "basis": report.basis or "",
                "approver": report.approver or default_approver_for(report.employee),
                "comment": report.comment or "",
            }
        )
        return context

    def build_file_name(self, report: BusinessTripReport, template_path: Path) -> str:
        person = _surname(report.employee.full_name)
        date_part = report.report_date.isoformat()
        city = slugify_file_part(report.trip_city, "Город")
        prefix = slugify_file_part(template_path.stem)
        return f"{prefix}_{person}_{city}_{date_part}.docx"


class RepresentativeExpenseBuilder(BaseReportBuilder):
    report_type = "representative_expenses"

    def build(self, report: RepresentativeExpenseReport) -> BuildResult:
        self.output_dir.mkdir(parents=True, exist_ok=True)
        output_path = unique_path(self.output_dir / self.build_file_name(report, Path("Смета_и_отчет_представительские.docx")))
        document = _build_representative_expense_document(report)
        document.save(str(output_path))
        return BuildResult(files=[output_path], warnings=[])

    def build_context(self, report: RepresentativeExpenseReport) -> dict[str, Any]:
        return _base_context(report.receipts, report.report_date, initiator=report.initiator) | {
            "event_date": format_date_ru(report.event_date),
            "place": report.place,
            "restaurant_name": report.restaurant_name,
            "counterparty": report.counterparty,
            "meeting_purpose": report.meeting_purpose,
            "participants_company": report.participants_company,
            "participants_counterparty": report.participants_counterparty,
            "participants_company_text": "; ".join(report.participants_company),
            "participants_counterparty_text": "; ".join(report.participants_counterparty),
            "meeting_result": report.meeting_result,
        }

    def build_file_name(self, report: RepresentativeExpenseReport, template_path: Path) -> str:
        counterparty = slugify_file_part(report.counterparty, "Контрагент")
        prefix = slugify_file_part(template_path.stem)
        return f"{prefix}_{counterparty}_{report.report_date.strftime('%d%m%Y')}.docx"


class GiftExpenseBuilder(BaseReportBuilder):
    report_type = "gifts"

    def build(self, report: GiftExpenseReport) -> BuildResult:
        self.output_dir.mkdir(parents=True, exist_ok=True)
        output_path = unique_path(self.output_dir / self.build_file_name(report, Path("Служебная_записка_подарки.docx")))
        document = _build_gift_expense_memo(report)
        document.save(str(output_path))
        return BuildResult(files=[output_path], warnings=[])

    def build_context(self, report: GiftExpenseReport) -> dict[str, Any]:
        return _base_context(report.receipts, report.report_date, initiator=report.initiator) | {
            "purchase_date": format_date_ru(report.purchase_date),
            "gift_name": report.gift_name,
            "gift_quantity": report.gift_quantity,
            "unit_price": format_rubles(report.unit_price),
            "unit_price_raw": report.unit_price,
            "gift_total": format_rubles(report.calculated_gift_amount),
            "recipients": report.recipients,
            "recipients_text": "; ".join(report.recipients),
            "counterparty": report.counterparty,
            "occasion": report.occasion,
            "purpose": report.purpose,
        }

    def build_file_name(self, report: GiftExpenseReport, template_path: Path) -> str:
        counterparty = slugify_file_part(report.counterparty, "Контрагент")
        prefix = slugify_file_part(template_path.stem)
        return f"{prefix}_{counterparty}_{report.report_date.strftime('%d%m%Y')}.docx"


def _build_gift_expense_memo(report: GiftExpenseReport) -> Document:
    document = Document()
    _setup_gift_document(document)

    for _ in range(2):
        document.add_paragraph()
    _add_right_paragraph(document, DEFAULT_SIGNATORY_POSITION)
    _add_right_paragraph(document, "ООО «ХУАСЮНЬ ГРУПП РУ»")
    _add_right_paragraph(document, DEFAULT_SIGNATORY_NAME_DATIVE)
    _add_right_paragraph(document, f"от {_position_from(report.initiator.position)}")
    _add_right_paragraph(document, _full_name_genitive(report.initiator.full_name))

    for _ in range(3):
        document.add_paragraph()
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Служебная записка о компенсации расходов")
    title_run.bold = True
    title_run.font.name = "Times New Roman"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    title_run.font.size = Pt(12)

    document.add_paragraph()
    total_amount = report.total_amount if report.receipts else report.calculated_gift_amount
    request = document.add_paragraph()
    request.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    request.add_run(
        "Прошу компенсировать расходы, понесенные с целью "
        f"{_gift_purpose_for_sentence(report.purpose)}, "
        f"в размере {format_rubles(total_amount).replace(' ₽', '')} "
        f"({_amount_words_integer(total_amount)} рублей"
    )
    kopecks = int((total_amount - int(total_amount)) * 100)
    if kopecks:
        request.add_run(f" {kopecks:02d} коп")
    request.add_run(") рублей.")

    document.add_paragraph()
    document.add_paragraph()
    document.add_paragraph("Приложения:")
    document.add_paragraph()
    attachments_numbering_id = _new_representative_numbering_id(document)
    for index, receipt in enumerate(report.receipts, start=1):
        _add_representative_numbered_paragraph(document, _gift_receipt_application(receipt, index), attachments_numbering_id)

    for _ in range(3):
        document.add_paragraph()
    _add_representative_signature_block(document, report.initiator)

    document.add_paragraph()
    document.add_paragraph(f"Дата {report.report_date.strftime('%d.%m.%Y')}")
    return document


def _gift_purpose_for_sentence(purpose: str) -> str:
    value = " ".join((purpose or "").split())
    if not value:
        return "поддержания деловых отношений"

    normalized = value.lower().replace("ё", "е")
    default_nominative = (
        "создание долгосрочных деловых отношений, укрепление связей с ключевыми клиентами "
        "и деловыми партнерами и формирование корпоративного имиджа и деловой репутации"
    )
    if normalized == default_nominative:
        return (
            "создания долгосрочных деловых отношений, укрепления связей с ключевыми клиентами "
            "и деловыми партнерами и формирования корпоративного имиджа и деловой репутации"
        )
    return value


def _setup_gift_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)


def _add_right_paragraph(document: Document, text: str):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)
    return paragraph


def _build_representative_expense_document(report: RepresentativeExpenseReport) -> Document:
    document = Document()
    _setup_representative_document(document)

    _add_representative_header(document, report.initiator)
    _add_blank_paragraphs(document, 3)
    _add_representative_title(document, "Смета на проведение переговоров")
    _add_blank_paragraphs(document, 1)
    _add_representative_paragraph(
        document,
        f"Прошу согласовать проведение переговоров с {_counterparty_phrase(report.counterparty)}.",
    )
    _add_blank_paragraphs(document, 1)
    _add_representative_paragraph(document, f"Участники переговоров: {_representative_participants_text(report)}")
    _add_blank_paragraphs(document, 1)
    _add_representative_paragraph(document, f"Дата проведения переговоров: {report.event_date.strftime('%d.%m.%Y')}")
    _add_blank_paragraphs(document, 1)
    _add_representative_paragraph(document, "Планируемые результаты деловой встречи:")
    planned_numbering_id = _new_representative_numbering_id(document)
    for line in _representative_planned_lines(report):
        _add_representative_numbered_paragraph(document, line, planned_numbering_id)
    _add_blank_paragraphs(document, 4)
    _add_representative_signature_block(document, report.initiator)
    _add_blank_paragraphs(document, 3)
    _add_representative_paragraph(document, f"Дата {report.report_date.strftime('%d.%m.%Y')}")

    document.add_page_break()

    _add_representative_header(document, report.initiator)
    _add_blank_paragraphs(document, 1)
    _add_representative_title(document, "Отчет о проведении переговоров")
    _add_blank_paragraphs(document, 1)
    _add_representative_paragraph(
        document,
        "Прошу возместить расходы на проведение переговоров в сумме "
        f"{_amount_in_rubles_text(report.total_amount)}.",
    )
    _add_blank_paragraphs(document, 1)
    _add_representative_paragraph(document, f"Участники переговоров: {_representative_participants_text(report)}")
    _add_blank_paragraphs(document, 1)
    _add_representative_paragraph(document, f"Дата проведения переговоров: {report.event_date.strftime('%d.%m.%Y')}")
    _add_blank_paragraphs(document, 1)
    _add_representative_paragraph(document, _representative_place_text(report))
    _add_blank_paragraphs(document, 1)
    _add_representative_paragraph(document, "Результаты деловой встречи:")
    result_numbering_id = _new_representative_numbering_id(document)
    for line in _representative_result_lines(report):
        _add_representative_numbered_paragraph(document, line, result_numbering_id)
    _add_blank_paragraphs(document, 1)
    _add_representative_paragraph(document, "Приложения:")
    attachments_numbering_id = _new_representative_numbering_id(document)
    for index, receipt in enumerate(report.receipts, start=1):
        _add_representative_numbered_paragraph(document, _representative_receipt_application(receipt, index), attachments_numbering_id)
    _add_blank_paragraphs(document, 3)
    _add_representative_signature_block(document, report.initiator)
    _add_blank_paragraphs(document, 3)
    _add_representative_paragraph(document, f"Дата {report.report_date.strftime('%d.%m.%Y')}")
    return document


def _setup_representative_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)


def _add_representative_header(document: Document, employee: Employee) -> None:
    for line in [
        DEFAULT_SIGNATORY_POSITION,
        "ООО «ХУАСЮНЬ ГРУПП РУ»",
        DEFAULT_SIGNATORY_NAME_DATIVE,
        f"от {_position_from(employee.position)}",
        _full_name_genitive(employee.full_name),
    ]:
        paragraph = _add_representative_paragraph(document, line)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _add_representative_title(document: Document, text: str):
    paragraph = _add_representative_paragraph(document, text, bold=True)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return paragraph


def _add_representative_paragraph(document: Document, text: str = "", *, bold: bool = False):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)
    return paragraph


def _add_representative_numbered_paragraph(document: Document, text: str, numbering_id: int):
    paragraph = _add_representative_paragraph(document, _capitalize_first_letter(text))
    paragraph.style = "List Number"
    _set_paragraph_numbering(paragraph, numbering_id)
    paragraph.paragraph_format.left_indent = Cm(1.0)
    paragraph.paragraph_format.first_line_indent = Cm(-0.4)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Cm(1.0))
    return paragraph


def _new_representative_numbering_id(document: Document) -> int:
    numbering = document.part.numbering_part.element
    number = numbering.add_num(7)
    level_override = number.add_lvlOverride(0)
    level_override.add_startOverride(1)
    return int(number.get(qn("w:numId")))


def _set_paragraph_numbering(paragraph, numbering_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.get_or_add_numPr()
    ilvl = num_pr.get_or_add_ilvl()
    ilvl.set(qn("w:val"), "0")
    num_id = num_pr.get_or_add_numId()
    num_id.set(qn("w:val"), str(numbering_id))


def _capitalize_first_letter(text: str) -> str:
    for index, char in enumerate(text):
        if char.isalpha():
            return f"{text[:index]}{char.upper()}{text[index + 1:]}"
    return text


def _add_blank_paragraphs(document: Document, count: int) -> None:
    for _ in range(count):
        _add_representative_paragraph(document)


def _representative_participants_text(report: RepresentativeExpenseReport) -> str:
    participants = []
    participants.extend(
        _participant_with_company(participant, report.counterparty)
        for participant in report.participants_counterparty
        if participant.strip()
    )
    participants.extend(
        _participant_with_company(participant, "Хуасюнь Групп Ру")
        for participant in report.participants_company
        if participant.strip()
    )
    return "; ".join(participants) or "участники деловой встречи"


def _participant_with_company(participant: str, company: str) -> str:
    value = participant.strip()
    if not value:
        return ""
    if "(" in value and ")" in value:
        before, after = value.split("(", 1)
        bracket_value = after.rsplit(")", 1)[0].strip()
        before = before.strip()
        if "," in before:
            name, position = (part.strip() for part in before.split(",", 1))
            return f"{name} ({position}, {bracket_value})"
        return value
    company = company.strip()
    if not company:
        return value
    if "," not in value:
        return f"{value} ({company})"
    name, position = (part.strip() for part in value.split(",", 1))
    return f"{name} ({position}, {company})"


def _representative_planned_lines(report: RepresentativeExpenseReport) -> list[str]:
    lines = _split_multiline_text(report.meeting_purpose)
    if lines:
        return lines
    return ["обсуждение текущих и перспективных проектов", "согласование дальнейшего порядка взаимодействия"]


def _representative_result_lines(report: RepresentativeExpenseReport) -> list[str]:
    lines = _split_multiline_text(report.meeting_result)
    if lines:
        return lines
    purpose_lines = _split_multiline_text(report.meeting_purpose)
    if purpose_lines:
        return purpose_lines
    return ["проведены переговоры и согласованы дальнейшие шаги по взаимодействию"]


def _split_multiline_text(value: str) -> list[str]:
    cleaned_lines = []
    for raw_line in value.replace(";", "\n").splitlines():
        line = raw_line.strip(" \t-•")
        if line:
            cleaned_lines.append(line)
    return cleaned_lines


def _counterparty_phrase(counterparty: str) -> str:
    value = counterparty.strip()
    if not value:
        return "контрагентом"
    lower = value.lower()
    if lower.startswith(("компан", "девелоп", "ооо", "ао", "пао", "зао", "оао", "гк", "ип", "«")):
        return value
    return f"компанией {value}"


def _representative_place_text(report: RepresentativeExpenseReport) -> str:
    receipt_places = _representative_receipt_places(report)
    if len(receipt_places) > 1:
        return f"Места переговоров: {'; '.join(receipt_places)}"
    if len(receipt_places) == 1 and (not report.restaurant_name.strip() or not report.place.strip()):
        return f"Место переговоров: {receipt_places[0]}"
    return f"Место переговоров: {_representative_place_item(report.restaurant_name, report.place)}"


def _representative_receipt_places(report: RepresentativeExpenseReport) -> list[str]:
    places = []
    seen = set()
    for receipt in report.receipts:
        if not receipt.seller and not receipt.address:
            continue
        place = _representative_place_item(receipt.seller or "", receipt.address or "")
        key = place.lower()
        if key in seen:
            continue
        seen.add(key)
        places.append(place)
    return places


def _representative_place_item(restaurant_name: str, place: str) -> str:
    restaurant = restaurant_name.strip()
    place = place.strip()
    if restaurant:
        lower = restaurant.lower()
        restaurant_text = restaurant if lower.startswith(("ресторан", "кафе", "бар", "кофейня")) else f"ресторан «{restaurant}»"
    else:
        restaurant_text = "место проведения переговоров"
    return restaurant_text + (f" ({place})" if place else "")


def _representative_receipt_application(receipt: Receipt, index: int) -> str:
    number = (
        receipt.fiscal_document_number
        or receipt.check_number
        or receipt.fiscal_number
        or Path(receipt.file_name).stem
        or str(index)
    )
    return f"Фискальный чек №{number}"


def _gift_receipt_application(receipt: Receipt, index: int) -> str:
    number = receipt.check_number or receipt.fiscal_number or receipt.fiscal_document_number or str(index)
    date_text = receipt.date.strftime("%d.%m.%Y") if receipt.date else ""
    return f"Кассовый чек №{number}" + (f" от {date_text}" if date_text else "")


def _add_representative_signature_block(document: Document, employee: Employee) -> None:
    rows = 5 if employee.manager_name else 2
    table = document.add_table(rows=rows, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = (Cm(9.1), Cm(4.6), Cm(4.2))
    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Pt(18)
        for index, cell in enumerate(row.cells):
            cell.width = widths[index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _clear_cell(cell)

    _set_cell_text(table.cell(0, 0), "Инициатор", bold=True)
    _set_cell_text(table.cell(1, 0), employee.position)
    _set_signature_line_cell(table.cell(1, 1))
    _set_cell_text(table.cell(1, 2), f"/{_signature_name(employee.full_name)}/")

    if not employee.manager_name:
        return

    table.rows[2].height = Pt(20)
    _set_cell_text(table.cell(3, 0), "Согласовано", bold=True)
    _set_cell_text(table.cell(4, 0), employee.manager_position or "")
    _set_signature_line_cell(table.cell(4, 1))
    _set_cell_text(table.cell(4, 2), f"/{_signature_name(employee.manager_name)}/")


def _clear_cell(cell) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1


def _set_cell_text(cell, text: str, *, bold: bool = False) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)


def _set_signature_line_cell(cell) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_cell_bottom_border(cell)


def _set_cell_bottom_border(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "0")
    bottom.set(qn("w:color"), "000000")


def _base_context(
    receipts: list[Receipt],
    report_date,
    employee: Employee | None = None,
    initiator: Employee | None = None,
) -> dict[str, Any]:
    empty_employee = _empty_employee()
    employee_context = _employee_context(employee or initiator) if (employee or initiator) else empty_employee
    initiator_context = _employee_context(initiator or employee) if (initiator or employee) else empty_employee
    total = sum((receipt.amount for receipt in receipts), start=Decimal("0"))
    by_type: dict[str, Any] = {}
    for receipt in receipts:
        by_type[receipt.expense_type] = by_type.get(receipt.expense_type, Decimal("0")) + receipt.amount
    return {
        "employee": employee_context,
        "initiator": initiator_context,
        "report_date": format_date_ru(report_date),
        "report_date_short": report_date.strftime("%d.%m.%Y"),
        "report_date_iso": report_date.isoformat(),
        "receipts": [_receipt_context(index, receipt) for index, receipt in enumerate(receipts, start=1)],
        "receipts_table": _receipts_table_text(receipts),
        "receipt_attachments_text": _receipt_attachments_text(receipts),
        "receipts_count": len(receipts),
        "total_amount": format_rubles(total),
        "total_amount_no_kopecks": format_rubles(total, with_kopecks=False),
        "total_amount_integer": str(int(total)),
        "total_amount_raw": total,
        "total_amount_words": amount_to_words_ru(total),
        "total_amount_words_integer": _amount_words_integer(total),
        "total_amount_integer_text": _amount_in_rubles_text(total),
        "expenses_by_type": {key: format_rubles(value) for key, value in by_type.items()},
        "taxi_compensation_purpose": _taxi_compensation_purpose(receipts),
        "taxi_trip_dates": _receipt_dates_range(receipts),
    }


def _employee_context(employee: Employee) -> dict[str, str]:
    data = employee.model_dump()
    context = {key: "" if value is None else str(value) for key, value in data.items()}
    context["default_signatory_name"] = context["default_signatory_name"] or DEFAULT_SIGNATORY_NAME
    context["default_signatory_name_dative"] = DEFAULT_SIGNATORY_NAME_DATIVE
    context["default_signatory_position"] = context["default_signatory_position"] or DEFAULT_SIGNATORY_POSITION
    context["full_name_genitive"] = _full_name_genitive(employee.full_name)
    context["position_from"] = _position_from(employee.position)
    context["signature_name"] = _signature_name(employee.full_name)
    context["manager_signature_name"] = _signature_name(employee.manager_name or "")
    context["approval_title"] = "Согласовано" if employee.manager_name else ""
    if not employee.manager_name:
        context["manager_position"] = "__NO_APPROVAL__"
        context["manager_signature_name"] = "__NO_APPROVAL__"
        context["approval_title"] = "__NO_APPROVAL__"
    return context


def _empty_employee() -> dict[str, str]:
    return {
        "id": "",
        "full_name": "",
        "full_name_genitive": "",
        "short_name": "",
        "position": "",
        "position_from": "",
        "department": "",
        "company": "",
        "phone": "",
        "email": "",
        "manager_name": "",
        "manager_position": "",
        "manager_signature_name": "",
        "approval_title": "",
        "default_signatory_name": "",
        "default_signatory_name_dative": "",
        "default_signatory_position": "",
        "signature_name": "",
    }


def _receipt_context(index: int, receipt: Receipt) -> dict[str, str]:
    return {
        "number": str(index),
        "file_name": receipt.file_name,
        "date": format_date_ru(receipt.date),
        "date_short": receipt.date.strftime("%d.%m.%Y") if receipt.date else "",
        "seller": receipt.seller or "",
        "address": receipt.address or "",
        "inn": receipt.inn or "",
        "amount": format_rubles(receipt.amount),
        "amount_no_kopecks": format_rubles(receipt.amount, with_kopecks=False),
        "amount_integer": str(int(receipt.amount)),
        "expense_type": receipt.expense_type,
        "comment": receipt.comment or "",
        "route": receipt.route or "",
        "fiscal_number": receipt.fiscal_number or "",
        "check_number": receipt.check_number or receipt.fiscal_number or str(index),
        "shift_number": receipt.shift_number or "",
        "kkt_number": receipt.kkt_number or "",
        "fiscal_document_number": receipt.fiscal_document_number or "",
        "fiscal_drive_number": receipt.fiscal_drive_number or "",
        "fiscal_sign": receipt.fiscal_sign or "",
        "payment_type": receipt.payment_type or "",
        "qr_raw": receipt.qr_raw or "",
    }


def _receipts_table_text(receipts: list[Receipt]) -> str:
    if not receipts:
        return "Чеки не добавлены"
    rows = []
    for index, receipt in enumerate(receipts, start=1):
        rows.append(
            f"{index}. {format_date_ru(receipt.date)} | {receipt.seller or '-'} | {receipt.address or '-'} | "
            f"{receipt.expense_type} | {format_rubles(receipt.amount)} | {receipt.comment or ''}"
        )
    return "\n".join(rows)


def _receipt_attachments_text(receipts: list[Receipt]) -> str:
    if not receipts:
        return ""
    rows = []
    for index, receipt in enumerate(receipts, start=1):
        number = receipt.check_number or receipt.fiscal_number or str(index)
        date_text = receipt.date.strftime("%d.%m.%Y") if receipt.date else ""
        rows.append(f"{index}) Чек №{number}" + (f" от {date_text}" if date_text else ""))
    return "\n".join(rows)


def _receipt_dates_range(receipts: list[Receipt]) -> str:
    dates = sorted(receipt.date for receipt in receipts if receipt.date)
    if not dates:
        return ""
    first = dates[0].strftime("%d.%m.%Y")
    last = dates[-1].strftime("%d.%m.%Y")
    return first if first == last else f"{first}-{last}"


def _taxi_compensation_purpose(receipts: list[Receipt]) -> str:
    parts = []
    for index, receipt in enumerate(receipts, start=1):
        number = receipt.check_number or receipt.fiscal_number or str(index)
        date_text = receipt.date.strftime("%d.%m.%Y") if receipt.date else ""
        route_or_comment = receipt.route or receipt.comment or "поездка на такси"
        suffix = f" (чек №{number}" + (f" от {date_text}" if date_text else "") + f" на сумму {format_rubles(receipt.amount)})"
        parts.append(f"{route_or_comment}{suffix}")
    return ", ".join(parts)


def _business_trip_taxi_compensation_purpose(
    receipts: list[Receipt],
    *,
    trip_city: str,
    trip_start_date: date,
    trip_end_date: date,
    trip_purpose: str = "",
) -> str:
    if not receipts:
        return ""
    airport_receipts = _airport_receipts(receipts)
    outbound_airport_receipt = airport_receipts[0] if airport_receipts else None
    return_airport_receipt = airport_receipts[-1] if len(airport_receipts) > 1 else None
    arrival_hotel_receipt = _arrival_hotel_receipt(receipts, outbound_airport_receipt, return_airport_receipt)
    auto_receipts_seen = 0
    parts = []
    sorted_receipts = sorted(
        receipts,
        key=lambda receipt: _taxi_purpose_sort_key(
            receipt,
            receipts,
            outbound_airport_receipt,
            return_airport_receipt,
            arrival_hotel_receipt,
        ),
    )
    for receipt in sorted_receipts:
        number = receipt.check_number or receipt.fiscal_number or str(receipts.index(receipt) + 1)
        date_text = receipt.date.strftime("%d.%m.%Y") if receipt.date else ""
        if receipt.route or receipt.comment:
            purpose = receipt.route or receipt.comment or ""
        else:
            is_outbound_airport = receipt is outbound_airport_receipt
            is_return_airport = receipt is return_airport_receipt
            is_arrival_hotel = receipt is arrival_hotel_receipt
            purpose = _automatic_taxi_purpose(
                receipt,
                is_outbound_airport=is_outbound_airport,
                is_return_airport=is_return_airport,
                is_arrival_hotel=is_arrival_hotel,
                trip_city=trip_city,
                trip_purpose=trip_purpose,
                sequence_index=auto_receipts_seen,
            )
            if not is_outbound_airport and not is_return_airport and not is_arrival_hotel:
                auto_receipts_seen += 1
        suffix = f" (чек №{number}" + (f" от {date_text}" if date_text else "") + f" на сумму {format_rubles(receipt.amount)})"
        parts.append(f"{purpose}{suffix}")
    return ", ".join(parts)


def _automatic_taxi_purpose(
    receipt: Receipt,
    *,
    is_outbound_airport: bool,
    is_return_airport: bool,
    is_arrival_hotel: bool,
    trip_city: str,
    trip_purpose: str,
    sequence_index: int,
) -> str:
    city_text = f" в г. {trip_city}" if trip_city else ""
    client = _city_client_name(trip_city, sequence_index)
    if is_outbound_airport:
        return "поездка в аэропорт Москвы для вылета в командировку"
    if is_return_airport:
        return "поездка из аэропорта Москвы после возвращения из командировки"
    if is_arrival_hotel:
        city_text = f" в г. {trip_city}" if trip_city else ""
        return f"поездка из аэропорта в отель{city_text}"
    purpose_specific = _purpose_based_taxi_route(trip_purpose, trip_city, sequence_index)
    if purpose_specific:
        return purpose_specific
    client_visit_purposes = [
        f"поездка на встречу с представителями {client} на строительном объекте{city_text}",
        f"поездка с объекта {client} на рабочую встречу по проекту{city_text}",
        f"поездка в гостиницу{city_text} после рабочей встречи с представителями {client}",
    ]
    return client_visit_purposes[sequence_index % len(client_visit_purposes)]


def _purpose_based_taxi_route(trip_purpose: str, trip_city: str, sequence_index: int) -> str:
    purpose = trip_purpose.lower().strip()
    if not purpose:
        return ""
    if sequence_index >= 2:
        return ""
    city_text = f" в г. {trip_city}" if trip_city else ""
    client = _city_client_name(trip_city, sequence_index)
    if any(keyword in purpose for keyword in ("выстав", "форум", "экспо", "конференц")):
        variants = [
            f"поездка на выставку{city_text}",
            f"поездка с выставки на рабочую встречу с представителями {client}{city_text}",
            f"поездка в гостиницу{city_text} после деловой программы выставки",
        ]
        return variants[sequence_index % len(variants)]
    if any(keyword in purpose for keyword in ("переговор", "встреч", "презентац")):
        variants = [
            f"поездка на переговоры с представителями {client}{city_text}",
            f"поездка с переговоров на строительный объект {client}{city_text}",
            f"поездка в гостиницу{city_text} после рабочей встречи с представителями {client}",
        ]
        return variants[sequence_index % len(variants)]
    if any(keyword in purpose for keyword in ("объект", "строй", "строит", "монтаж", "обслед", "аудит")):
        variants = [
            f"поездка на строительный объект {client}{city_text}",
            f"поездка с объекта {client} на техническое совещание{city_text}",
            f"поездка в гостиницу{city_text} после технического совещания по объекту {client}",
        ]
        return variants[sequence_index % len(variants)]
    return ""


def _airport_receipts(receipts: list[Receipt]) -> list[Receipt]:
    if len(receipts) < 3:
        return sorted(receipts, key=lambda receipt: receipt.amount, reverse=True)[:1]
    expensive_receipts = sorted(receipts, key=lambda receipt: receipt.amount, reverse=True)[:2]
    return sorted(
        expensive_receipts,
        key=lambda receipt: (receipt.date or date.max, receipts.index(receipt)),
    )


def _compact_business_trip_taxi_memo(path: Path, *, receipts_count: int) -> None:
    document = Document(str(path))
    for section in document.sections:
        section.top_margin = Inches(0.45)
        section.bottom_margin = Inches(0.45)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    for paragraph in document.paragraphs:
        if paragraph.text.strip() == DEFAULT_SIGNATORY_NAME:
            _replace_paragraph_text(paragraph, DEFAULT_SIGNATORY_NAME_DATIVE)
        _normalize_rubles_suffix(paragraph)
        _compact_paragraph(paragraph)
        if (
            paragraph.text.startswith("Прошу ")
            or paragraph.text.startswith("Даты поездок:")
            or paragraph.text.startswith("Цель поездок:")
            or paragraph.text.startswith("Цель поездки:")
        ):
            paragraph.paragraph_format.space_after = Pt(6)
        if paragraph.text.startswith("Цель поездок:") or paragraph.text.startswith("Цель поездки:"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    _limit_empty_paragraphs_before_first_table(document, keep_count=4)
    for table in document.tables:
        _remove_empty_approval_rows(table)
        for row_index, row in enumerate(table.rows):
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            row.height = Pt(10 if row_index == 2 and receipts_count >= 5 else 16 if receipts_count >= 5 else 20)
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _compact_paragraph(paragraph)

    document.save(str(path))


def _normalize_rubles_suffix(paragraph) -> None:
    normalized_text = re.sub(r"\)\s+рубл(?:я|ей)\s*\.?", ") рублей.", paragraph.text)
    if normalized_text == paragraph.text:
        return
    _replace_paragraph_text(paragraph, normalized_text)


def _remove_empty_approval_rows(table) -> None:
    if len(table.rows) < 5:
        return
    approval_title = table.cell(3, 0).text.strip()
    approval_position = table.cell(4, 0).text.strip()
    approval_signature = table.cell(4, 2).text.strip()
    no_approval_values = {"", "//", "__NO_APPROVAL__", "/__NO_APPROVAL__/"}
    if approval_title not in no_approval_values or approval_position not in no_approval_values or approval_signature not in no_approval_values:
        return
    for row_index in (4, 3, 2):
        table._tbl.remove(table.rows[row_index]._tr)


def _compact_paragraph(paragraph) -> None:
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0


def _replace_paragraph_text(paragraph, text: str) -> None:
    alignment = paragraph.alignment
    style = paragraph.style
    first_run = paragraph.runs[0] if paragraph.runs else None
    font_name = first_run.font.name if first_run else "Times New Roman"
    font_size = first_run.font.size if first_run else Pt(12)
    is_bold = bool(first_run.bold) if first_run else False
    paragraph.clear()
    paragraph.style = style
    paragraph.alignment = alignment
    run = paragraph.add_run(text)
    run.bold = is_bold
    run.font.name = font_name or "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name or "Times New Roman")
    run.font.size = font_size or Pt(12)


def _limit_empty_paragraphs_before_first_table(document, *, keep_count: int) -> None:
    body = document.element.body
    children = list(body)
    first_table_index = next((index for index, child in enumerate(children) if child.tag.endswith("}tbl")), None)
    if first_table_index is None:
        return

    empty_paragraphs = []
    for child in reversed(children[:first_table_index]):
        if not child.tag.endswith("}p"):
            break
        text = "".join(node.text or "" for node in child.iter() if node.tag.endswith("}t")).strip()
        has_break = any(node.tag.endswith("}br") for node in child.iter())
        if text or has_break:
            break
        empty_paragraphs.append(child)

    for paragraph_element in empty_paragraphs[keep_count:]:
        body.remove(paragraph_element)


def _arrival_hotel_receipt(
    receipts: list[Receipt],
    outbound_airport_receipt: Receipt | None,
    return_airport_receipt: Receipt | None,
) -> Receipt | None:
    if not outbound_airport_receipt or not outbound_airport_receipt.date:
        return None
    same_day_receipts = [
        receipt
        for receipt in receipts
        if receipt is not outbound_airport_receipt
        and receipt is not return_airport_receipt
        and receipt.date == outbound_airport_receipt.date
        and not receipt.route
        and not receipt.comment
    ]
    if not same_day_receipts:
        return None
    return max(same_day_receipts, key=lambda receipt: (receipt.amount, -receipts.index(receipt)))


def _taxi_purpose_sort_key(
    receipt: Receipt,
    receipts: list[Receipt],
    outbound_airport_receipt: Receipt | None,
    return_airport_receipt: Receipt | None,
    arrival_hotel_receipt: Receipt | None,
) -> tuple[date, int, int]:
    receipt_date = receipt.date or date.max
    rank = 1
    if receipt is outbound_airport_receipt:
        rank = 0
    elif receipt is arrival_hotel_receipt:
        rank = 1
    elif receipt is return_airport_receipt:
        rank = 3
    elif outbound_airport_receipt and receipt.date == outbound_airport_receipt.date:
        rank = 2
    elif return_airport_receipt and receipt.date == return_airport_receipt.date:
        rank = 0
    return receipt_date, rank, receipts.index(receipt)


def _city_client_name(trip_city: str, sequence_index: int = 0) -> str:
    clients = {
        "Екатеринбург": ["ООО «СЛМ»", "ГК «Урал Девелопмент»", "ООО «СтройИнвест-Урал»", "АО «Высота Проект»"],
        "Москва": ["АО «Мосинжпроект»", "ГК «Прайм Девелопмент»", "ООО «СтройКонтур»", "АО «Городские Проекты»"],
        "Санкт-Петербург": ["ГК «Эталон»", "ООО «Северный Девелопмент»", "АО «Лахта Строй»", "ООО «Петербург Проект»"],
        "Казань": ["ГК «Унистрой»", "ООО «Казань Девелопмент»", "АО «Волга Строй»", "ООО «СитиПроект Казань»"],
        "Новосибирск": ["ГК «Стрижи»", "ООО «Сибирь Девелопмент»", "АО «Обь Строй»", "ООО «Новострой Проект»"],
        "Сочи": ["ГК «Метрополис»", "ООО «Юг Девелопмент»", "АО «Черномор Строй»", "ООО «Курорт Проект»"],
        "Самара": ["ГК «Новый Дон»", "ООО «Самара Девелопмент»", "АО «Волга Проект»", "ООО «СтройГрад Самара»"],
        "Нижний Новгород": ["ГК «Столица Нижний»", "ООО «Ока Девелопмент»", "АО «НН Строй»", "ООО «Верхневолжский Проект»"],
        "Краснодар": ["ГК «ССК»", "ООО «Кубань Девелопмент»", "АО «Южный Строй»", "ООО «Краснодар Проект»"],
        "Уфа": ["ГК «Жилстройинвест»", "ООО «Уфа Девелопмент»", "АО «БашСтройПроект»", "ООО «Белая Река Строй»"],
    }
    city_clients = clients.get(trip_city.strip(), ["ООО «СЛМ»", "ООО «Регион Девелопмент»", "АО «ГородСтрой»"])
    return city_clients[sequence_index % len(city_clients)]


def _amount_words_integer(amount) -> str:
    try:
        from num2words import num2words

        return num2words(int(amount), lang="ru")
    except Exception:
        return amount_to_words_ru(amount).split(" руб", 1)[0]


def _amount_in_rubles_text(amount) -> str:
    return f"{int(amount)} ({_amount_words_integer(amount)}) рублей"


def _signature_name(full_name: str) -> str:
    parts = full_name.split()
    if not parts:
        return ""
    surname = parts[0]
    initials = "".join(f"{part[0]}." for part in parts[1:] if part)
    return f"{surname} {initials}".strip()


def _full_name_genitive(full_name: str) -> str:
    known_names = {
        "Баранова Гиляна Басанговна": "Барановой Гиляны Басанговны",
        "Другалев Александр Александрович": "Другалева Александра Александровича",
        "Конопельнюк Антон Петрович": "Конопельнюка Антона Петровича",
        "Платонов Антон Александрович": "Платонова Антона Александровича",
        "Зимин Сергей Александрович": "Зимина Сергея Александровича",
        "Попов Леонид Николаевич": "Попова Леонида Николаевича",
    }
    return known_names.get(full_name, full_name)


def _position_from(position: str) -> str:
    known_positions = {
        "Руководитель направления продаж": "руководителя направления продаж",
        "Менеджер по продажам": "менеджера по продажам",
        "Директор по развитию проектных продаж": "директора по развитию проектных продаж",
        "Руководитель проекта по развитию": "руководителя проекта по развитию",
        "Руководитель проекта": "руководителя проекта",
        "Менеджер проектов": "менеджера проектов",
        "Менеджер по работе с ключевыми клиентами": "менеджера по работе с ключевыми клиентами",
    }
    return known_positions.get(position, position[:1].lower() + position[1:] if position else "")


def _surname(full_name: str) -> str:
    return slugify_file_part(full_name.split()[0], "Сотрудник")
