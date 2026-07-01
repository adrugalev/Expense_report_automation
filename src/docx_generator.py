from __future__ import annotations

import re
import zipfile
from pathlib import Path
from typing import Any

from docx import Document

try:
    from docxtpl import DocxTemplate
except Exception:  # pragma: no cover - fallback is covered instead
    DocxTemplate = None

from .utils import unique_path


PLACEHOLDER_RE = re.compile(r"{{\s*([a-zA-Z_][\w.]*?)\s*}}")


class DocxGenerator:
    def render(self, template_path: Path, context: dict[str, Any], output_path: Path) -> Path:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path = unique_path(output_path)
        missing = self.find_unfilled_placeholders(template_path, context)
        if DocxTemplate is not None:
            template = DocxTemplate(str(template_path))
            template.render(context)
            template.save(output_path)
        else:
            self._render_simple(template_path, context, output_path)
        if missing:
            warning_path = output_path.with_suffix(".warnings.txt")
            warning_path.write_text(
                "Не заполнены плейсхолдеры:\n" + "\n".join(sorted(missing)),
                encoding="utf-8",
            )
        return output_path

    def find_unfilled_placeholders(self, template_path: Path, context: dict[str, Any]) -> set[str]:
        text = extract_docx_text(template_path)
        placeholders = set(PLACEHOLDER_RE.findall(text))
        return {name for name in placeholders if _resolve_context_path(context, name) in (None, "")}

    def _render_simple(self, template_path: Path, context: dict[str, Any], output_path: Path) -> None:
        document = Document(str(template_path))
        for paragraph in document.paragraphs:
            _replace_in_paragraph(paragraph, context)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        _replace_in_paragraph(paragraph, context)
        document.save(str(output_path))


def extract_docx_text(path: Path) -> str:
    with zipfile.ZipFile(path) as docx:
        xml_parts = [
            name
            for name in docx.namelist()
            if name.startswith("word/") and name.endswith(".xml")
        ]
        return "\n".join(docx.read(name).decode("utf-8", errors="ignore") for name in xml_parts)


def _replace_in_paragraph(paragraph, context: dict[str, Any]) -> None:
    full_text = "".join(run.text for run in paragraph.runs)
    replaced = PLACEHOLDER_RE.sub(lambda match: str(_resolve_context_path(context, match.group(1)) or ""), full_text)
    if replaced == full_text:
        return
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = replaced
    else:
        paragraph.add_run(replaced)


def _resolve_context_path(context: dict[str, Any], path: str) -> Any:
    current: Any = context
    for part in path.split("."):
        if isinstance(current, dict):
            current = current.get(part)
        else:
            current = getattr(current, part, None)
        if current is None:
            return None
    return current

