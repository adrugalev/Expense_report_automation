from datetime import date
from decimal import Decimal
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from src.docx_generator import extract_docx_text
from src.models import BusinessTripReport, Employee, GiftExpenseReport, Receipt, RepresentativeExpenseReport
from src.report_builders import (
    BusinessTripBuilder,
    GiftExpenseBuilder,
    RepresentativeExpenseBuilder,
    _amount_in_rubles_text,
    _gift_purpose_for_sentence,
)
from src.template_manager import TemplateManager


def test_business_trip_context_contains_employee_and_totals(tmp_path):
    employee = Employee(
        full_name="Иванов Иван",
        short_name="И. Иванов",
        position="Менеджер",
        department="Отдел",
        company="ООО Тест",
    )
    report = BusinessTripReport(
        employee=employee,
        trip_city="Москва",
        trip_start_date=date(2026, 6, 20),
        trip_end_date=date(2026, 6, 25),
        purpose="Встреча",
        receipts=[Receipt(file_name="check.jpg", date=date(2026, 6, 21), amount=Decimal("100"), expense_type="такси")],
        report_date=date(2026, 6, 26),
    )
    builder = BusinessTripBuilder(TemplateManager(tmp_path / "templates"), tmp_path / "output")
    context = builder.build_context(report)
    assert context["employee"]["full_name"] == "Иванов Иван"
    assert context["total_amount"] == "100,00 ₽"
    assert context["receipts"][0]["expense_type"] == "такси"


def test_business_trip_uses_only_taxi_memo_template(tmp_path):
    templates_dir = tmp_path / "templates"
    manager = TemplateManager(templates_dir)
    manager.ensure_default_templates("business_trip")

    old_template = templates_dir / "business_trip" / "Реестр_чеков_командировка.docx"
    document = Document()
    document.add_paragraph("Старый шаблон")
    document.save(old_template)

    assert [path.name for path in manager.templates_for("business_trip")] == ["Служебная_записка_такси.docx"]


def test_representative_expenses_use_single_combined_template(tmp_path):
    templates_dir = tmp_path / "templates"
    manager = TemplateManager(templates_dir)
    manager.ensure_default_templates("representative_expenses")

    old_template = templates_dir / "representative_expenses" / "Отчет_представительские.docx"
    document = Document()
    document.add_paragraph("Старый шаблон")
    document.save(old_template)

    assert [path.name for path in manager.templates_for("representative_expenses")] == ["Смета_и_отчет_представительские.docx"]


def test_business_trip_context_contains_taxi_memo_defaults(tmp_path):
    employee = Employee(
        full_name="Другалев Александр Александрович",
        short_name="А.А. Другалев",
        position="Руководитель направления продаж",
        company="ООО «Хуасюнь Групп Ру»",
    )
    report = BusinessTripReport(
        employee=employee,
        trip_city="Екатеринбург",
        trip_start_date=date(2026, 6, 20),
        trip_end_date=date(2026, 6, 25),
        purpose="Встреча",
        receipts=[Receipt(file_name="check.jpg", date=date(2026, 6, 21), amount=Decimal("100"), expense_type="такси")],
        report_date=date(2026, 6, 26),
    )

    context = BusinessTripBuilder(TemplateManager(tmp_path / "templates"), tmp_path / "output").build_context(report)

    assert context["employee"]["default_signatory_position"] == "Генеральному директору"
    assert context["employee"]["default_signatory_name"] == "Хуан Голян"
    assert context["employee"]["default_signatory_name_dative"] == "Хуан Голяну"
    assert context["employee"]["position_from"] == "руководителя направления продаж"


def test_business_trip_taxi_purpose_uses_realistic_routes(tmp_path):
    employee = Employee(
        full_name="Баранова Гиляна Басанговна",
        short_name="Г.Б. Баранова",
        position="Менеджер по продажам",
        company="ООО «Хуасюнь Групп Ру»",
    )
    report = BusinessTripReport(
        employee=employee,
        trip_city="Екатеринбург",
        trip_start_date=date(2025, 9, 30),
        trip_end_date=date(2025, 10, 3),
        purpose="Встреча",
        receipts=[
            Receipt(file_name="596.pdf", date=date(2025, 10, 2), amount=Decimal("292"), expense_type="такси", check_number="596"),
            Receipt(file_name="422.pdf", date=date(2025, 10, 2), amount=Decimal("381"), expense_type="такси", check_number="422"),
            Receipt(file_name="292.pdf", date=date(2025, 9, 30), amount=Decimal("2596"), expense_type="такси", check_number="292"),
        ],
        report_date=date(2025, 10, 3),
    )

    context = BusinessTripBuilder(TemplateManager(tmp_path / "templates"), tmp_path / "output").build_context(report)
    purpose = context["taxi_compensation_purpose"]

    assert "поездка в аэропорт Москвы для вылета в командировку (чек №292 от 30.09.2025 на сумму 2 596,00 ₽)" in purpose
    assert "поездка на переговоры с представителями ООО «СЛМ» в г. Екатеринбург (чек №596 от 02.10.2025 на сумму 292,00 ₽)" in purpose
    assert "поездка из аэропорта Москвы после возвращения из командировки (чек №422 от 02.10.2025 на сумму 381,00 ₽)" in purpose


def test_business_trip_taxi_purpose_rotates_city_clients(tmp_path):
    employee = Employee(
        full_name="Баранова Гиляна Басанговна",
        short_name="Г.Б. Баранова",
        position="Менеджер по продажам",
        company="ООО «Хуасюнь Групп Ру»",
    )
    report = BusinessTripReport(
        employee=employee,
        trip_city="Екатеринбург",
        trip_start_date=date(2025, 9, 30),
        trip_end_date=date(2025, 10, 3),
        purpose="Встреча",
        receipts=[
            Receipt(file_name="596.pdf", date=date(2025, 10, 2), amount=Decimal("292"), expense_type="такси", check_number="596"),
            Receipt(file_name="422.pdf", date=date(2025, 10, 2), amount=Decimal("381"), expense_type="такси", check_number="422"),
            Receipt(file_name="412.pdf", date=date(2025, 10, 2), amount=Decimal("435"), expense_type="такси", check_number="412"),
            Receipt(file_name="292.pdf", date=date(2025, 9, 30), amount=Decimal("2596"), expense_type="такси", check_number="292"),
        ],
        report_date=date(2025, 10, 3),
    )

    context = BusinessTripBuilder(TemplateManager(tmp_path / "templates"), tmp_path / "output").build_context(report)
    purpose = context["taxi_compensation_purpose"]

    assert "представителями ООО «СЛМ»" in purpose
    assert "строительный объект ГК «Урал Девелопмент»" in purpose


def test_business_trip_taxi_purpose_uses_trip_purpose_for_exhibition(tmp_path):
    employee = Employee(
        full_name="Баранова Гиляна Басанговна",
        short_name="Г.Б. Баранова",
        position="Менеджер по продажам",
        company="ООО «Хуасюнь Групп Ру»",
    )
    report = BusinessTripReport(
        employee=employee,
        trip_city="Екатеринбург",
        trip_start_date=date(2025, 9, 30),
        trip_end_date=date(2025, 10, 3),
        purpose="посещение выставки",
        receipts=[
            Receipt(file_name="596.pdf", date=date(2025, 10, 2), amount=Decimal("292"), expense_type="такси", check_number="596"),
            Receipt(file_name="292.pdf", date=date(2025, 9, 30), amount=Decimal("2596"), expense_type="такси", check_number="292"),
        ],
        report_date=date(2025, 10, 3),
    )

    context = BusinessTripBuilder(TemplateManager(tmp_path / "templates"), tmp_path / "output").build_context(report)

    assert "поездка на выставку в г. Екатеринбург (чек №596 от 02.10.2025 на сумму 292,00 ₽)" in context["taxi_compensation_purpose"]


def test_business_trip_taxi_purpose_uses_trip_purpose_only_for_first_local_trips(tmp_path):
    employee = Employee(
        full_name="Баранова Гиляна Басановна",
        short_name="Г.Б. Баранова",
        position="Менеджер по продажам",
        company="ООО «Хуасюнь Групп Ру»",
        manager_name="Другалев Александр Александрович",
        manager_position="Руководитель направления продаж",
    )
    report = BusinessTripReport(
        employee=employee,
        trip_city="Екатеринбург",
        trip_start_date=date(2025, 9, 30),
        trip_end_date=date(2025, 10, 3),
        purpose="посещение выставки",
        receipts=[
            Receipt(file_name="292.pdf", date=date(2025, 9, 30), amount=Decimal("2596"), expense_type="такси", check_number="292"),
            Receipt(file_name="596.pdf", date=date(2025, 10, 2), amount=Decimal("292"), expense_type="такси", check_number="596"),
            Receipt(file_name="422.pdf", date=date(2025, 10, 2), amount=Decimal("381"), expense_type="такси", check_number="422"),
            Receipt(file_name="412.pdf", date=date(2025, 10, 2), amount=Decimal("435"), expense_type="такси", check_number="412"),
            Receipt(file_name="777.pdf", date=date(2025, 10, 3), amount=Decimal("3000"), expense_type="такси", check_number="777"),
        ],
        report_date=date(2025, 10, 3),
    )

    context = BusinessTripBuilder(TemplateManager(tmp_path / "templates"), tmp_path / "output").build_context(report)
    purpose = context["taxi_compensation_purpose"]

    assert purpose.count("выстав") == 2
    assert "поездка в гостиницу в г. Екатеринбург после рабочей встречи с представителями ООО «СтройИнвест-Урал»" in purpose
    assert "с представителями ООО «СтройИнвест-Урал» в гостиницу" not in purpose


def test_business_trip_taxi_purpose_sends_same_day_arrival_to_hotel(tmp_path):
    employee = Employee(
        full_name="Баранова Гиляна Басановна",
        short_name="Г.Б. Баранова",
        position="Менеджер по продажам",
        company="ООО «Хуасюнь Групп Ру»",
    )
    report = BusinessTripReport(
        employee=employee,
        trip_city="Екатеринбург",
        trip_start_date=date(2025, 9, 30),
        trip_end_date=date(2025, 10, 3),
        purpose="посещение выставки",
        receipts=[
            Receipt(file_name="89.pdf", date=date(2025, 9, 30), amount=Decimal("994"), expense_type="такси", check_number="89"),
            Receipt(file_name="292.pdf", date=date(2025, 9, 30), amount=Decimal("2596"), expense_type="такси", check_number="292"),
            Receipt(file_name="747.pdf", date=date(2025, 10, 2), amount=Decimal("502"), expense_type="такси", check_number="747"),
            Receipt(file_name="596.pdf", date=date(2025, 10, 2), amount=Decimal("292"), expense_type="такси", check_number="596"),
            Receipt(file_name="422.pdf", date=date(2025, 10, 2), amount=Decimal("381"), expense_type="такси", check_number="422"),
            Receipt(file_name="412.pdf", date=date(2025, 10, 2), amount=Decimal("435"), expense_type="такси", check_number="412"),
            Receipt(file_name="210.pdf", date=date(2025, 10, 3), amount=Decimal("4224"), expense_type="такси", check_number="210"),
        ],
        report_date=date(2025, 10, 3),
    )

    context = BusinessTripBuilder(TemplateManager(tmp_path / "templates"), tmp_path / "output").build_context(report)
    purpose = context["taxi_compensation_purpose"]

    assert context["taxi_compensation_purpose_label"] == "Цель поездок"
    assert purpose.index("чек №292") < purpose.index("чек №89")
    assert "поездка из аэропорта в отель в г. Екатеринбург (чек №89 от 30.09.2025 на сумму 994,00 ₽)" in purpose
    assert purpose.count("выстав") == 2


def test_business_trip_taxi_memo_renders_receipt_list_and_approval(tmp_path):
    employee = Employee(
        full_name="Баранова Гиляна Басанговна",
        short_name="Г.Б. Баранова",
        position="Менеджер по продажам",
        company="ООО «Хуасюнь Групп Ру»",
        manager_name="Другалев Александр Александрович",
        manager_position="Руководитель направления продаж",
    )
    report = BusinessTripReport(
        employee=employee,
        trip_city="Екатеринбург",
        trip_start_date=date(2025, 10, 2),
        trip_end_date=date(2025, 10, 2),
        purpose="Встреча",
        receipts=[
            Receipt(file_name="422.pdf", date=date(2025, 10, 2), amount=Decimal("381"), expense_type="такси", check_number="422"),
            Receipt(file_name="412.pdf", date=date(2025, 10, 2), amount=Decimal("412"), expense_type="такси", check_number="412"),
        ],
        report_date=date(2025, 10, 3),
    )

    result = BusinessTripBuilder(TemplateManager(Path("templates")), tmp_path / "output").build(report)
    text = extract_docx_text(result.files[0])
    generated = Document(result.files[0])
    attachments_paragraph = generated.paragraphs[15]
    signature_table = generated.tables[0]

    assert generated.paragraphs[4].text == "Хуан Голяну"
    assert generated.paragraphs[6].text.strip() == "Барановой Гиляны Басанговны"
    assert attachments_paragraph._p.pPr is None or attachments_paragraph._p.pPr.numPr is None
    assert "1) Чек №422 от 02.10.2025" in text
    assert "2) Чек №412 от 02.10.2025" in text
    assert "чек №422 от 02.10.2025 на сумму 381,00 ₽" in text
    assert "чек №412 от 02.10.2025 на сумму 412,00 ₽" in text
    assert "1) Чек №422 от 02.10.2025 на сумму" not in text
    assert signature_table.cell(0, 0).text == "Инициатор"
    assert signature_table.cell(1, 0).text == "Менеджер по продажам"
    assert signature_table.cell(1, 1).text == ""
    assert _cell_bottom_border(signature_table.cell(1, 1)) == "single"
    assert signature_table.cell(1, 2).text == "/Баранова Г.Б./"
    assert signature_table.cell(2, 0).text == ""
    assert signature_table.rows[2].height.pt >= 16
    assert signature_table.cell(3, 0).text == "Согласовано"
    assert len(signature_table.rows) == 5
    assert signature_table.cell(4, 0).text == "Руководитель направления продаж"
    assert signature_table.cell(4, 1).text == ""
    assert _cell_bottom_border(signature_table.cell(4, 1)) == "single"
    assert signature_table.cell(4, 2).text == "/Другалев А.А./"


def test_business_trip_taxi_memo_uses_rubley_suffix_for_amount_words(tmp_path):
    employee = Employee(
        full_name="Баранова Гиляна Басанговна",
        short_name="Г.Б. Баранова",
        position="Менеджер по продажам",
        company="ООО «Хуасюнь Групп Ру»",
        manager_name="Другалев Александр Александрович",
        manager_position="Руководитель направления продаж",
    )
    report = BusinessTripReport(
        employee=employee,
        trip_city="Екатеринбург",
        trip_start_date=date(2025, 10, 2),
        trip_end_date=date(2025, 10, 2),
        purpose="Встреча",
        receipts=[Receipt(file_name="taxi.pdf", date=date(2025, 10, 2), amount=Decimal("8627"), expense_type="такси", check_number="8627")],
        report_date=date(2025, 10, 3),
    )

    result = BusinessTripBuilder(TemplateManager(Path("templates")), tmp_path / "output").build(report)
    text = extract_docx_text(result.files[0])
    generated = Document(result.files[0])
    request_paragraph = next(paragraph for paragraph in generated.paragraphs if paragraph.text.startswith("Прошу компенсировать"))

    assert "8627 (восемь тысяч шестьсот двадцать семь) рублей." in text
    assert ") рубля" not in text
    assert max(run.font.size.pt for run in request_paragraph.runs if run.font.size) == 12


def test_business_trip_taxi_memo_omits_empty_approval_block(tmp_path):
    employee = Employee(
        full_name="Другалев Александр Александрович",
        short_name="А.А. Другалев",
        position="Руководитель направления продаж",
        company="ООО «Хуасюнь Групп Ру»",
    )
    report = BusinessTripReport(
        employee=employee,
        trip_city="Екатеринбург",
        trip_start_date=date(2025, 10, 2),
        trip_end_date=date(2025, 10, 2),
        purpose="Встреча",
        receipts=[
            Receipt(file_name="422.pdf", date=date(2025, 10, 2), amount=Decimal("381"), expense_type="такси", check_number="422"),
        ],
        report_date=date(2025, 10, 3),
    )

    result = BusinessTripBuilder(TemplateManager(Path("templates")), tmp_path / "output").build(report)
    generated = Document(result.files[0])
    signature_table = generated.tables[0]
    table_text = "\n".join(cell.text for row in signature_table.rows for cell in row.cells)

    assert result.warnings == []
    assert len(signature_table.rows) == 2
    assert signature_table.cell(0, 0).text == "Инициатор"
    assert signature_table.cell(1, 0).text == "Руководитель направления продаж"
    assert signature_table.cell(1, 2).text == "/Другалев А.А./"
    assert "Согласовано" not in table_text
    assert "//" not in table_text


def test_business_trip_taxi_memo_uses_dative_director_for_any_employee(tmp_path):
    employee = Employee(
        full_name="Попов Леонид Николаевич",
        position="Менеджер по работе с ключевыми клиентами",
        company="ООО «Хуасюнь Групп Ру»",
        manager_name="Другалев Александр Александрович",
        manager_position="Руководитель направления продаж",
    )
    report = BusinessTripReport(
        employee=employee,
        trip_city="Екатеринбург",
        trip_start_date=date(2025, 10, 2),
        trip_end_date=date(2025, 10, 2),
        purpose="Встреча",
        receipts=[
            Receipt(file_name="422.pdf", date=date(2025, 10, 2), amount=Decimal("381"), expense_type="такси", check_number="422"),
        ],
        report_date=date(2025, 10, 3),
    )

    result = BusinessTripBuilder(TemplateManager(Path("templates")), tmp_path / "output").build(report)
    generated = Document(result.files[0])

    assert generated.paragraphs[4].text == "Хуан Голяну"
    assert "Хуан Голян\n" not in "\n".join(paragraph.text for paragraph in generated.paragraphs[:8])


def test_gift_expense_builder_generates_single_memo_like_example(tmp_path):
    initiator = Employee(
        full_name="Другалев Александр Александрович",
        position="Руководитель направления продаж",
        company="ООО «Хуасюнь Групп Ру»",
    )
    report = GiftExpenseReport(
        initiator=initiator,
        purchase_date=date(2026, 6, 18),
        gift_name="подарочная продукция",
        gift_quantity=1,
        unit_price=Decimal("19520.96"),
        recipients=[],
        counterparty="Antteq",
        occasion="",
        purpose=(
            "создания долгосрочных деловых отношений, укрепления связей с ключевыми клиентами "
            "и деловыми партнерами и для формирования корпоративного имиджа и деловой репутации"
        ),
        receipts=[
            Receipt(
                file_name="gift.pdf",
                date=date(2026, 6, 18),
                amount=Decimal("19520.96"),
                expense_type="подарки",
                check_number="18724",
            )
        ],
        report_date=date(2026, 6, 19),
    )

    result = GiftExpenseBuilder(TemplateManager(Path("templates")), tmp_path / "output").build(report)
    generated = Document(result.files[0])
    text = "\n".join(paragraph.text for paragraph in generated.paragraphs)

    assert len(result.files) == 1
    assert generated.paragraphs[4].text == "Хуан Голяну"
    title = next(paragraph for paragraph in generated.paragraphs if paragraph.text == "Служебная записка о компенсации расходов")
    request = next(paragraph for paragraph in generated.paragraphs if paragraph.text.startswith("Прошу компенсировать"))
    attachment = next(paragraph for paragraph in generated.paragraphs if paragraph.text == "Кассовый чек №18724 от 18.06.2026")
    signature_table = generated.tables[0]
    assert title.runs[0].bold is True
    assert request.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
    assert attachment.style.name == "List Number"
    assert signature_table.cell(0, 0).text == "Инициатор"
    assert signature_table.cell(1, 0).text == "Руководитель направления продаж"
    assert len(signature_table.rows) == 2
    assert "Служебная записка о компенсации расходов" in text
    assert "19 520,96" in text
    assert "Кассовый чек №18724 от 18.06.2026" in text
    assert "Дата 19.06.2026" in text


def test_gift_purpose_accepts_nominative_default_for_form():
    purpose = (
        "создание долгосрочных деловых отношений, укрепление связей с ключевыми клиентами "
        "и деловыми партнерами и формирование корпоративного имиджа и деловой репутации"
    )

    assert _gift_purpose_for_sentence(purpose) == (
        "создания долгосрочных деловых отношений, укрепления связей с ключевыми клиентами "
        "и деловыми партнерами и формирования корпоративного имиджа и деловой репутации"
    )


def test_gift_expense_builder_adds_approval_for_non_drugalev(tmp_path):
    initiator = Employee(
        full_name="Баранова Гиляна Басанговна",
        position="Менеджер по продажам",
        company="ООО «Хуасюнь Групп Ру»",
        manager_name="Другалев Александр Александрович",
        manager_position="Руководитель направления продаж",
    )
    report = GiftExpenseReport(
        initiator=initiator,
        purchase_date=date(2026, 6, 18),
        gift_name="подарочная продукция",
        gift_quantity=1,
        unit_price=Decimal("3508"),
        recipients=[],
        counterparty="Antteq",
        occasion="",
        purpose="создания долгосрочных деловых отношений",
        receipts=[
            Receipt(file_name="gift.pdf", date=date(2026, 6, 18), amount=Decimal("3508"), expense_type="подарки", check_number="73859")
        ],
        report_date=date(2026, 6, 19),
    )

    result = GiftExpenseBuilder(TemplateManager(Path("templates")), tmp_path / "output").build(report)
    generated = Document(result.files[0])
    signature_table = generated.tables[0]

    assert len(signature_table.rows) == 5
    assert signature_table.cell(0, 0).text == "Инициатор"
    assert signature_table.cell(3, 0).text == "Согласовано"
    assert signature_table.cell(4, 0).text == "Руководитель направления продаж"
    assert signature_table.cell(4, 2).text == "/Другалев А.А./"


def test_representative_expense_builder_generates_combined_estimate_and_report(tmp_path):
    initiator = Employee(
        full_name="Баранова Гиляна Басанговна",
        short_name="Г.Б. Баранова",
        position="Менеджер по продажам",
        company="ООО «Хуасюнь Групп Ру»",
        manager_name="Другалев Александр Александрович",
        manager_position="Руководитель направления продаж",
    )
    report = RepresentativeExpenseReport(
        initiator=initiator,
        event_date=date(2026, 3, 5),
        place="Москва, Трубная улица, 18",
        restaurant_name="Smoke BBQ",
        counterparty="Upside Development",
        meeting_purpose="Обсуждение перспективных проектов",
        participants_company=["Баранова Гиляна Басанговна", "Другалев Александр Александрович"],
        participants_counterparty=["Иванов Иван, коммерческий директор (Upside Development)", "Петров Петр, главный инженер"],
        meeting_result="обсуждены перспективные проекты\nсогласован порядок дальнейшего взаимодействия",
        receipts=[
            Receipt(
                file_name="check.pdf",
                date=date(2026, 3, 5),
                amount=Decimal("19109"),
                expense_type="ресторан",
                fiscal_document_number="47280",
            )
        ],
        report_date=date(2026, 3, 6),
    )

    result = RepresentativeExpenseBuilder(TemplateManager(tmp_path / "templates"), tmp_path / "output").build(report)
    generated = Document(result.files[0])
    text = "\n".join(paragraph.text for paragraph in generated.paragraphs)

    assert result.warnings == []
    assert len(result.files) == 1
    assert result.files[0].name == "Смета_и_отчет_представительские_Upside_Development_06032026.docx"
    assert "Смета на проведение переговоров" in text
    assert "Отчет о проведении переговоров" in text
    assert text.count("Хуан Голяну") == 2
    assert "\nХуан Голян\n" not in f"\n{text}\n"
    assert "Прошу согласовать проведение переговоров с компанией Upside Development." in text
    assert "Прошу возместить расходы на проведение переговоров в сумме 19109" in text
    assert "девятнадцать тысяч сто девять" in text
    assert "Место переговоров: ресторан «Smoke BBQ» (Москва, Трубная улица, 18)" in text
    assert "Фискальный чек №47280" in text
    assert "Баранова Гиляна Басанговна (Хуасюнь Групп Ру)" in text
    assert "Другалев Александр Александрович (Хуасюнь Групп Ру)" in text
    assert "Иванов Иван (коммерческий директор, Upside Development)" in text
    assert "Петров Петр (главный инженер, Upside Development)" in text
    assert len(generated.tables) == 2
    assert generated.tables[0].cell(0, 0).text == "Инициатор"
    assert generated.tables[0].cell(3, 0).text == "Согласовано"
    assert _next_paragraph_text(generated, "Прошу согласовать") == ""
    assert _next_paragraph_text(generated, "Участники переговоров") == ""
    assert _next_paragraph_text(generated, "Дата проведения переговоров") == ""
    assert _next_paragraph_text(generated, "Прошу возместить") == ""
    assert _next_paragraph_text(generated, "Участники переговоров", occurrence=2) == ""
    assert _next_paragraph_text(generated, "Дата проведения переговоров", occurrence=2) == ""
    assert _next_paragraph_text(generated, "Место переговоров") == ""
    assert _next_paragraph_text(generated, "Согласован порядок дальнейшего взаимодействия") == ""
    assert next(paragraph for paragraph in generated.paragraphs if paragraph.text == "Обсуждение перспективных проектов").style.name == "List Number"
    assert next(paragraph for paragraph in generated.paragraphs if paragraph.text == "Согласован порядок дальнейшего взаимодействия").style.name == "List Number"
    assert next(paragraph for paragraph in generated.paragraphs if paragraph.text == "Фискальный чек №47280").style.name == "List Number"
    planned_num_id = _paragraph_num_id(next(paragraph for paragraph in generated.paragraphs if paragraph.text == "Обсуждение перспективных проектов"))
    result_num_id = _paragraph_num_id(next(paragraph for paragraph in generated.paragraphs if paragraph.text == "Согласован порядок дальнейшего взаимодействия"))
    attachment_num_id = _paragraph_num_id(next(paragraph for paragraph in generated.paragraphs if paragraph.text == "Фискальный чек №47280"))
    assert len({planned_num_id, result_num_id, attachment_num_id}) == 3


def test_representative_amount_text_always_uses_rubley_suffix():
    assert _amount_in_rubles_text(Decimal("8627")) == "8627 (восемь тысяч шестьсот двадцать семь) рублей"


def test_representative_expense_builder_lists_all_receipt_places_in_combined_report(tmp_path):
    initiator = Employee(
        full_name="Баранова Гиляна Басанговна",
        short_name="Г.Б. Баранова",
        position="Менеджер по продажам",
        company="ООО «Хуасюнь Групп Ру»",
        manager_name="Другалев Александр Александрович",
        manager_position="Руководитель направления продаж",
    )
    report = RepresentativeExpenseReport(
        initiator=initiator,
        event_date=date(2026, 6, 30),
        place="Общее место",
        restaurant_name="Общий ресторан",
        counterparty="Coldy",
        meeting_purpose="Обсуждение проектов",
        participants_company=["Баранова Гиляна Басанговна"],
        participants_counterparty=["Иванов Иван, коммерческий директор"],
        meeting_result="согласован порядок дальнейшего взаимодействия",
        receipts=[
            Receipt(
                file_name="first.pdf",
                date=date(2026, 6, 30),
                seller="Корчма",
                address="г. Москва, ул. Садовая-Кудринская, д. 3А",
                amount=Decimal("21990"),
                expense_type="ресторан",
                fiscal_document_number="17419203",
            ),
            Receipt(
                file_name="second.pdf",
                date=date(2026, 6, 30),
                seller="Mr Hot Рамен",
                address="г. Москва, наб. Пресненская, д. 10",
                amount=Decimal("3697"),
                expense_type="ресторан",
                fiscal_document_number="24071",
            ),
        ],
        report_date=date(2026, 6, 30),
    )

    result = RepresentativeExpenseBuilder(TemplateManager(tmp_path / "templates"), tmp_path / "output").build(report)
    generated = Document(result.files[0])
    text = "\n".join(paragraph.text for paragraph in generated.paragraphs)

    assert (
        "Места переговоров: ресторан «Корчма» (г. Москва, ул. Садовая-Кудринская, д. 3А); "
        "ресторан «Mr Hot Рамен» (г. Москва, наб. Пресненская, д. 10)"
    ) in text
    assert "Фискальный чек №17419203" in text
    assert "Фискальный чек №24071" in text


def test_representative_expense_builder_omits_approval_for_drugalev(tmp_path):
    initiator = Employee(
        full_name="Другалев Александр Александрович",
        short_name="А.А. Другалев",
        position="Руководитель направления продаж",
        company="ООО «Хуасюнь Групп Ру»",
    )
    report = RepresentativeExpenseReport(
        initiator=initiator,
        event_date=date(2026, 3, 5),
        place="Москва",
        restaurant_name="Smoke BBQ",
        counterparty="EPSS",
        meeting_purpose="Переговоры",
        participants_company=["Другалев Александр Александрович"],
        participants_counterparty=[],
        meeting_result="проведены переговоры",
        receipts=[Receipt(file_name="check.pdf", date=date(2026, 3, 5), amount=Decimal("1000"), expense_type="ресторан")],
        report_date=date(2026, 3, 6),
    )

    result = RepresentativeExpenseBuilder(TemplateManager(tmp_path / "templates"), tmp_path / "output").build(report)
    generated = Document(result.files[0])
    table_text = "\n".join(cell.text for table in generated.tables for row in table.rows for cell in row.cells)

    assert all(len(table.rows) == 2 for table in generated.tables)
    assert "Согласовано" not in table_text


def test_business_trip_taxi_memo_compacts_long_report(tmp_path):
    employee = Employee(
        full_name="Баранова Гиляна Басановна",
        short_name="Г.Б. Баранова",
        position="Менеджер по продажам",
        company="ООО «Хуасюнь Групп Ру»",
        manager_name="Другалев Александр Александрович",
        manager_position="Руководитель направления продаж",
    )
    report = BusinessTripReport(
        employee=employee,
        trip_city="Екатеринбург",
        trip_start_date=date(2025, 9, 30),
        trip_end_date=date(2025, 10, 3),
        purpose="посещение выставки",
        receipts=[
            Receipt(file_name="89.pdf", date=date(2025, 9, 30), amount=Decimal("994"), expense_type="такси", check_number="89"),
            Receipt(file_name="292.pdf", date=date(2025, 9, 30), amount=Decimal("2596"), expense_type="такси", check_number="292"),
            Receipt(file_name="747.pdf", date=date(2025, 10, 2), amount=Decimal("502"), expense_type="такси", check_number="747"),
            Receipt(file_name="596.pdf", date=date(2025, 10, 2), amount=Decimal("292"), expense_type="такси", check_number="596"),
            Receipt(file_name="422.pdf", date=date(2025, 10, 2), amount=Decimal("381"), expense_type="такси", check_number="422"),
            Receipt(file_name="412.pdf", date=date(2025, 10, 2), amount=Decimal("435"), expense_type="такси", check_number="412"),
            Receipt(file_name="210.pdf", date=date(2025, 10, 3), amount=Decimal("4224"), expense_type="такси", check_number="210"),
        ],
        report_date=date(2025, 10, 3),
    )

    result = BusinessTripBuilder(TemplateManager(Path("templates")), tmp_path / "output").build(report)
    generated = Document(result.files[0])
    request_paragraph = next(paragraph for paragraph in generated.paragraphs if paragraph.text.startswith("Прошу "))
    dates_paragraph = next(paragraph for paragraph in generated.paragraphs if paragraph.text.startswith("Даты поездок"))
    purpose_paragraph = next(paragraph for paragraph in generated.paragraphs if paragraph.text.startswith("Цель поездок"))

    assert generated.sections[0].top_margin.inches <= 0.45
    assert _empty_paragraphs_before_first_table(generated) == 4
    assert max(run.font.size.pt for run in purpose_paragraph.runs if run.font.size) == 12
    assert request_paragraph.paragraph_format.space_after.pt == 6
    assert dates_paragraph.paragraph_format.space_after.pt == 6
    assert purpose_paragraph.paragraph_format.space_after.pt == 6
    assert purpose_paragraph.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
    assert generated.tables[0].rows[2].height.pt <= 10


def _empty_paragraphs_before_first_table(document: Document) -> int:
    children = list(document.element.body)
    first_table_index = next(index for index, child in enumerate(children) if child.tag.endswith("}tbl"))
    count = 0
    for child in reversed(children[:first_table_index]):
        if not child.tag.endswith("}p"):
            break
        text = "".join(node.text or "" for node in child.iter() if node.tag.endswith("}t")).strip()
        has_break = any(node.tag.endswith("}br") for node in child.iter())
        if text or has_break:
            break
        count += 1
    return count


def _next_paragraph_text(document: Document, startswith: str, occurrence: int = 1) -> str:
    paragraphs = list(document.paragraphs)
    matches = [index for index, paragraph in enumerate(paragraphs) if paragraph.text.lstrip().startswith(startswith)]
    index = matches[occurrence - 1]
    return paragraphs[index + 1].text


def _cell_bottom_border(cell) -> str | None:
    borders = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcBorders")
    if borders is None:
        return None
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        return None
    return bottom.get(qn("w:val"))


def _paragraph_num_id(paragraph) -> str | None:
    num_pr = paragraph._p.pPr.numPr if paragraph._p.pPr is not None else None
    if num_pr is None or num_pr.numId is None:
        return None
    return num_pr.numId.val
